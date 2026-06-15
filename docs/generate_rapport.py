from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"C:\Users\danad\Documents\Claude\Projects\projetmobile\docs\rapport_projet.docx"

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── Styles ────────────────────────────────────────────────────────────────────
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)

h1 = doc.styles['Heading 1']
h1.font.name  = 'Calibri'
h1.font.size  = Pt(16)
h1.font.bold  = True
h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

h2 = doc.styles['Heading 2']
h2.font.name  = 'Calibri'
h2.font.size  = Pt(13)
h2.font.bold  = True
h2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

h3 = doc.styles['Heading 3']
h3.font.name  = 'Calibri'
h3.font.size  = Pt(11)
h3.font.bold  = True
h3.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── Helper functions ──────────────────────────────────────────────────────────
def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_para(text, bold=False, italic=False, size=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if align:
        p.alignment = align
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5)
    return p

def add_numbered(text):
    return doc.add_paragraph(text, style='List Number')

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    return p

def page_break():
    doc.add_page_break()

def add_table_styled(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue fill
        tc   = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  '2E74B5')
        tcPr.append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx+1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()   # spacing after table

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RAPPORT DE PROJET")
run.bold      = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FleetTracking")
run.bold      = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Application Mobile de Gestion de Flotte de Véhicules")
run.font.size  = Pt(14)
run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Projet Universitaire  |  Développement Mobile & Backend")
run.font.size = Pt(11)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Technologies : Android (Java)  •  Spring Boot  •  PostgreSQL  •  Docker")
run.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2025 – 2026")
run.font.size = Pt(12)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE DES MATIÈRES (manuelle)
# ══════════════════════════════════════════════════════════════════════════════
add_h1("Table des matières")

toc_entries = [
    ("1.", "Introduction / Présentation du projet"),
    ("2.", "Architecture technique"),
    ("3.", "Modèle de données"),
    ("4.", "Fonctionnalités principales"),
    ("5.", "Système d'entretiens et notifications"),
    ("6.", "Interface utilisateur"),
    ("7.", "Déploiement et configuration"),
    ("8.", "Conclusion / Perspectives"),
]
for num, title in toc_entries:
    p = doc.add_paragraph()
    run1 = p.add_run(f"{num}  ")
    run1.bold = True
    p.add_run(title)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
add_h1("1. Introduction / Présentation du projet")

add_para(
    "Dans un contexte où la mobilité urbaine et la logistique de transport occupent une place centrale "
    "dans l'économie moderne, la gestion efficace d'une flotte de véhicules représente un enjeu "
    "stratégique majeur pour les entreprises. Les solutions traditionnelles, souvent papier ou tableur, "
    "ne permettent plus de répondre aux exigences de réactivité, de traçabilité et de sécurité imposées "
    "par les standards actuels."
)

add_para(
    "Le projet FleetTracking a été conçu et développé dans le cadre d'un projet universitaire de "
    "développement mobile et backend. Il s'agit d'une application Android complète, connectée à un "
    "serveur Spring Boot, permettant à une organisation de gérer sa flotte de véhicules en temps réel, "
    "de suivre ses chauffeurs, de centraliser la gestion des entretiens et des incidents, et d'assurer "
    "une traçabilité complète des trajets effectués."
)

add_h2("1.1 Objectifs du projet")

add_para("FleetTracking vise à atteindre les objectifs suivants :")
add_bullet("Fournir une interface mobile intuitive pour deux types d'utilisateurs : les administrateurs et les chauffeurs.")
add_bullet("Permettre le suivi GPS en temps réel des véhicules depuis l'application chauffeur.")
add_bullet("Centraliser la gestion des entretiens préventifs (vidanges, contrôles techniques) avec des alertes automatiques.")
add_bullet("Offrir un système de déclaration et de suivi des incidents.")
add_bullet("Générer des statistiques et une vue cartographique pour les administrateurs.")
add_bullet("Garantir la fiabilité des données grâce à une synchronisation backend et un cache local SQLite.")

add_h2("1.2 Périmètre fonctionnel")

add_para(
    "L'application couvre l'ensemble du cycle de vie d'une flotte : de l'enregistrement des véhicules "
    "et des chauffeurs jusqu'à la clôture des trajets, en passant par la planification des entretiens "
    "et la gestion des incidents. Deux profils utilisateurs distincts sont pris en charge, chacun "
    "disposant d'un espace dédié et sécurisé."
)

add_h2("1.3 Contexte technologique")

add_para(
    "Le choix des technologies a été guidé par les contraintes universitaires et les standards de "
    "l'industrie. Android (Java) a été retenu pour le client mobile, en raison de sa maturité et de "
    "sa large adoption. Spring Boot assure la robustesse du backend REST, tandis que PostgreSQL, "
    "déployé via Docker, garantit la persistance et l'intégrité des données. Google Maps SDK offre "
    "les capacités de géolocalisation et de cartographie."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  2. ARCHITECTURE TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════
add_h1("2. Architecture technique")

add_para(
    "FleetTracking repose sur une architecture client-serveur de type REST. L'application Android "
    "joue le rôle de client léger, délégant la logique métier au backend Spring Boot. "
    "La communication s'effectue sur le réseau local via HTTP/JSON."
)

add_h2("2.1 Vue d'ensemble de l'architecture")

add_code_block(
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│                    APPLICATION ANDROID (Client)                  │\n"
    "│                                                                  │\n"
    "│  ┌──────────────┐  ┌──────────────┐  ┌──────────────────────┐  │\n"
    "│  │  UI Layer    │  │  Repository  │  │  Services/Managers   │  │\n"
    "│  │  Activities  │◄─┤  (Retrofit + │  │  TripManager         │  │\n"
    "│  │  Fragments   │  │   SQLite)    │  │  WorkManager         │  │\n"
    "│  └──────────────┘  └──────┬───────┘  │  NotificationSvc     │  │\n"
    "│                           │           └──────────────────────┘  │\n"
    "└───────────────────────────┼─────────────────────────────────────┘\n"
    "                            │ HTTP/JSON (Retrofit + OkHttp)\n"
    "                            │ Réseau local (Wi-Fi / LAN)\n"
    "                            ▼\n"
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│                    BACKEND SPRING BOOT                          │\n"
    "│                                                                  │\n"
    "│  ┌──────────────┐  ┌──────────────┐  ┌──────────────────────┐  │\n"
    "│  │  Controllers  │  │  Services    │  │  Repositories        │  │\n"
    "│  │  REST API     │─►│  Logique     │─►│  Spring Data JPA     │  │\n"
    "│  │  (JSON)       │  │  métier      │  │                      │  │\n"
    "│  └──────────────┘  └──────────────┘  └──────────┬───────────┘  │\n"
    "└──────────────────────────────────────────────────┼──────────────┘\n"
    "                                                    │ JDBC\n"
    "                                                    ▼\n"
    "┌─────────────────────────────────────────────────────────────────┐\n"
    "│              PostgreSQL (Docker Container)                       │\n"
    "│              Port 5432  |  Volume persistant                     │\n"
    "└─────────────────────────────────────────────────────────────────┘"
)

add_h2("2.2 Couche Android")

add_h3("2.2.1 Structure des packages")
add_para("L'application Android est organisée selon une architecture en couches :")
add_bullet("com.fleettracking.app : activités principales (Login)")
add_bullet("com.fleettracking.app.admin : espace administrateur (activités + fragments)")
add_bullet("com.fleettracking.app.chauffeur : espace chauffeur (activités + fragments)")
add_bullet("com.fleettracking.app.data : couche données (Repository, modèles API)")
add_bullet("com.fleettracking.app.util : utilitaires (TripManager, ImageUtils, FleetConfig)")

add_h3("2.2.2 Bibliothèques clés")
add_table_styled(
    ["Bibliothèque", "Version", "Rôle"],
    [
        ["Retrofit 2", "2.9+", "Client HTTP REST"],
        ["OkHttp 3", "4.x", "Couche transport HTTP, logging"],
        ["Gson", "2.10+", "Sérialisation/désérialisation JSON"],
        ["WorkManager", "2.8+", "Tâches de fond périodiques (notifications)"],
        ["Google Maps SDK", "18.x", "Cartographie et géolocalisation"],
        ["Glide", "4.x", "Chargement et mise en cache d'images"],
        ["SharedPreferences", "Android API", "Persistance session utilisateur"],
        ["SQLite (Room)", "Android API", "Cache hors-ligne"],
    ]
)

add_h2("2.3 Backend Spring Boot")

add_para(
    "Le backend expose une API REST consommée exclusivement par l'application Android. "
    "Il est développé en Java avec Spring Boot et suit le pattern MVC :"
)
add_bullet("Controllers : reçoivent les requêtes HTTP, délèguent aux services, retournent du JSON.")
add_bullet("Services : contiennent la logique métier (calcul des prochains entretiens, validation des données).")
add_bullet("Repositories : interfaces Spring Data JPA générant automatiquement les requêtes SQL.")
add_bullet("Entities : classes Java mappées aux tables PostgreSQL via annotations JPA/Hibernate.")

add_h2("2.4 Communication réseau")

add_para(
    "La communication entre l'application Android et le serveur Spring Boot utilise le protocole HTTP "
    "sur le réseau local. L'adresse IP du serveur est configurable via la classe FleetConfig. "
    "Retrofit gère la transformation automatique des objets Java en JSON et vice-versa grâce à "
    "un convertisseur Gson. Un intercepteur OkHttp ajoute les en-têtes d'authentification "
    "(token JWT stocké dans SharedPreferences) à chaque requête."
)

add_para("Exemple de configuration Retrofit :")
add_code_block(
    "// FleetConfig.java\n"
    "public static final String BASE_URL = \"http://192.168.1.X:8080/api/\";\n\n"
    "// RetrofitClient.java\n"
    "Retrofit retrofit = new Retrofit.Builder()\n"
    "    .baseUrl(FleetConfig.BASE_URL)\n"
    "    .addConverterFactory(GsonConverterFactory.create())\n"
    "    .client(okHttpClient)\n"
    "    .build();"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  3. MODÈLE DE DONNÉES
# ══════════════════════════════════════════════════════════════════════════════
add_h1("3. Modèle de données")

add_para(
    "Le modèle de données de FleetTracking est structuré autour de six entités principales, "
    "reflétant les objets du domaine métier de la gestion de flotte. Ces entités sont "
    "persistées dans PostgreSQL via Spring Data JPA et mappées côté Android sous forme "
    "de classes POJO utilisées par Retrofit/Gson."
)

add_h2("3.1 Entité Vehicule")
add_para("Représente un véhicule de la flotte.")
add_table_styled(
    ["Champ", "Type", "Description"],
    [
        ["id", "Long", "Identifiant unique (auto-généré)"],
        ["marque", "String", "Marque du véhicule (ex. Renault, Peugeot)"],
        ["modele", "String", "Modèle du véhicule"],
        ["immatriculation", "String", "Numéro de plaque d'immatriculation"],
        ["annee", "Integer", "Année de mise en service"],
        ["kilometrage", "Double", "Kilométrage actuel du compteur"],
        ["statut", "Enum", "Disponible / En mission / Indisponible"],
        ["consommation", "Double", "Consommation moyenne (L/100 km)"],
        ["photo", "String", "Photo encodée en Base64"],
        ["vidangeCibleKm", "Double", "Kilométrage cible de la prochaine vidange"],
        ["controleTechniqueDate", "String", "Date du prochain contrôle technique (yyyy-MM-dd)"],
        ["conducteurId", "Long", "Référence au chauffeur assigné (FK)"],
        ["lat / lng", "Double", "Coordonnées GPS de dernière position connue"],
        ["vitesse", "Double", "Vitesse instantanée (km/h)"],
    ]
)

add_h2("3.2 Entité Chauffeur")
add_para("Représente un chauffeur enregistré dans le système.")
add_table_styled(
    ["Champ", "Type", "Description"],
    [
        ["id", "Long", "Identifiant unique"],
        ["nom", "String", "Nom complet du chauffeur"],
        ["telephone", "String", "Numéro de téléphone"],
        ["email", "String", "Adresse e-mail"],
        ["permis", "String", "Numéro de permis de conduire"],
        ["login", "String", "Identifiant de connexion"],
        ["motDePasse", "String", "Mot de passe (hashé côté backend)"],
        ["statut", "String", "Statut actif/inactif"],
        ["photo", "String", "Photo de profil encodée en Base64"],
    ]
)

add_h2("3.3 Entité Entretien")
add_para(
    "Représente une opération de maintenance planifiée ou effectuée. "
    "Deux types de déclenchement sont supportés : kilométrique et temporel."
)
add_table_styled(
    ["Champ", "Type", "Description"],
    [
        ["id", "Long", "Identifiant unique"],
        ["vehiculeId", "Long", "Référence au véhicule (FK)"],
        ["vehiculeNom", "String", "Nom du véhicule (dénormalisé pour affichage)"],
        ["immatriculation", "String", "Immatriculation (dénormalisée)"],
        ["type", "String", "Type d'entretien (Vidange, Contrôle Technique, etc.)"],
        ["estKmBase", "Boolean", "true = déclenchement km, false = déclenchement date"],
        ["cibleKm", "Double", "Kilométrage cible si estKmBase=true"],
        ["cibleDate", "String", "Date cible si estKmBase=false (yyyy-MM-dd)"],
        ["intervalKm", "Double", "Intervalle de récurrence en km"],
        ["intervalJours", "Integer", "Intervalle de récurrence en jours"],
        ["statut", "Enum", "aVenir / effectue"],
        ["dateEffectuee", "String", "Date de réalisation effective"],
        ["retardJours", "Integer", "Nombre de jours de retard calculé"],
    ]
)

add_h2("3.4 Entité Incident")
add_para("Représente un incident déclaré par un chauffeur ou un administrateur.")
add_table_styled(
    ["Champ", "Type", "Description"],
    [
        ["id", "Long", "Identifiant unique"],
        ["vehiculeId", "Long", "Référence au véhicule concerné (FK)"],
        ["vehiculeNom", "String", "Nom du véhicule (dénormalisé)"],
        ["type", "String", "Type d'incident (panne mécanique, accident, pneu crevé, etc.)"],
        ["description", "String", "Description textuelle de l'incident"],
        ["date", "String", "Date et heure de l'incident"],
        ["statut", "String", "En cours / Résolu"],
        ["photos", "List<String>", "Liste de photos encodées en Base64 (multi-images)"],
        ["lat / lng", "Double", "Coordonnées GPS du lieu de l'incident"],
    ]
)

add_h2("3.5 Entité Trajet")
add_para("Représente un trajet effectué par un chauffeur avec un véhicule.")
add_table_styled(
    ["Champ", "Type", "Description"],
    [
        ["id", "Long", "Identifiant unique"],
        ["vehiculeId", "Long", "Référence au véhicule (FK)"],
        ["vehiculeNom", "String", "Nom du véhicule (dénormalisé)"],
        ["chauffeurId", "Long", "Référence au chauffeur (FK)"],
        ["date", "String", "Date du trajet"],
        ["heureDepart", "String", "Heure de départ (format HH:mm:ss)"],
        ["heureArrivee", "String", "Heure d'arrivée (format HH:mm:ss)"],
        ["distanceKm", "Double", "Distance parcourue en kilomètres"],
        ["duree", "Long", "Durée en secondes"],
        ["vitesseMoyenne", "Double", "Vitesse moyenne calculée (km/h)"],
        ["consommationEstimee", "Double", "Consommation estimée (L) basée sur la distance"],
    ]
)

add_h2("3.6 Entité User")
add_para(
    "Entité de sécurité gérant l'authentification. Liée à un Chauffeur ou à un profil "
    "administrateur. Stocke le rôle (ADMIN / CHAUFFEUR) et les credentials de connexion. "
    "Le backend valide les credentials et retourne un token de session."
)

add_h2("3.7 Diagramme de relations")
add_code_block(
    "Vehicule ──< Entretien    (1 véhicule peut avoir plusieurs entretiens)\n"
    "Vehicule ──< Incident     (1 véhicule peut avoir plusieurs incidents)\n"
    "Vehicule ──< Trajet       (1 véhicule peut avoir plusieurs trajets)\n"
    "Chauffeur ──< Trajet      (1 chauffeur peut effectuer plusieurs trajets)\n"
    "Chauffeur ──○ Vehicule    (0 ou 1 véhicule assigné à un chauffeur)\n"
    "User ──○ Chauffeur        (0 ou 1 compte chauffeur lié à un utilisateur)"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  4. FONCTIONNALITÉS PRINCIPALES
# ══════════════════════════════════════════════════════════════════════════════
add_h1("4. Fonctionnalités principales")

add_h2("4.1 Authentification")
add_para(
    "L'écran de connexion (LoginActivity) est le point d'entrée unique de l'application. "
    "L'utilisateur saisit son identifiant et son mot de passe, puis sélectionne son rôle "
    "(Admin ou Chauffeur). Le système envoie une requête POST à l'endpoint /api/auth/login "
    "du backend Spring Boot. En cas de succès, le token de session et les informations "
    "utilisateur (id, nom, rôle, véhicule assigné pour les chauffeurs) sont persistés dans "
    "SharedPreferences et l'utilisateur est redirigé vers son espace dédié."
)
add_bullet("Gestion des erreurs : affichage d'un message en cas d'identifiants incorrects.")
add_bullet("Déconnexion : disponible depuis le profil, efface le token et redirige vers Login.")
add_bullet("Persistance de session : si un token valide est présent au démarrage, l'utilisateur est "
           "directement redirigé sans ressaisir ses credentials.")

add_h2("4.2 Espace Administrateur")

add_h3("4.2.1 Tableau de bord")
add_para(
    "L'écran d'accueil administrateur affiche des cartes de statistiques synthétiques : "
    "nombre total de véhicules, nombre de chauffeurs actifs, nombre d'incidents en cours, "
    "et nombre d'entretiens à venir. Ces données sont récupérées via des appels API "
    "au démarrage de l'activité."
)

add_h3("4.2.2 Gestion des véhicules")
add_para("L'administrateur dispose d'un écran listant tous les véhicules de la flotte avec :")
add_bullet("Barre de recherche par marque, modèle ou immatriculation.")
add_bullet("Filtre par statut (Disponible / En mission / Indisponible).")
add_bullet("Accès aux boutons Incidents et Entretiens pour chaque véhicule directement depuis la liste.")
add_bullet("Bouton FAB pour ajouter un nouveau véhicule.")
add_para("L'écran de détail d'un véhicule (VehiculeDetailsActivity) permet :")
add_bullet("La modification de la marque, du modèle, de la plaque d'immatriculation, de l'année, "
           "du kilométrage et de la consommation.")
add_bullet("La mise à jour du statut via un Spinner (liste déroulante).")
add_bullet("L'assignation d'un chauffeur via un Spinner dynamique (liste des chauffeurs disponibles).")
add_bullet("La saisie du kilométrage cible de la prochaine vidange.")
add_bullet("La saisie de la date du prochain contrôle technique via un DatePicker natif Android.")
add_bullet("L'ajout ou la modification de la photo du véhicule (galerie ou appareil photo).")
add_bullet("L'affichage de la position du véhicule sur une carte Google Maps.")

add_h3("4.2.3 Gestion des chauffeurs")
add_para("La liste des chauffeurs offre une recherche par nom ou email. "
         "L'écran de détail d'un chauffeur (ChauffeurDetailsActivity) permet :")
add_bullet("Modification du nom, téléphone, email, numéro de permis, identifiant de connexion et mot de passe.")
add_bullet("Assignation d'un véhicule via un Spinner.")
add_bullet("Gestion de la photo de profil.")

add_h3("4.2.4 Gestion des incidents")
add_para(
    "L'écran des incidents (IncidentsActivity) présente les incidents dans trois onglets "
    "(TabLayout + ViewPager) : Tous, En cours, Résolus. Chaque incident affiche le véhicule "
    "concerné, le type, la date et le statut. L'administrateur peut consulter le détail d'un "
    "incident (IncidentDetailsActivity), voir les photos jointes, et marquer l'incident comme résolu."
)

add_h3("4.2.5 Gestion des entretiens")
add_para(
    "L'écran des entretiens propose deux onglets : À venir et Historique. "
    "Chaque carte d'entretien affiche un badge coloré indiquant l'urgence :"
)
add_bullet("Vert : entretien non imminent (> 500 km ou > 7 jours).")
add_bullet("Orange : entretien imminent (≤ 500 km ou ≤ 7 jours).")
add_bullet("Rouge : entretien en retard (kilométrage ou date dépassés).")
add_para("Actions disponibles :")
add_bullet("Marquer comme effectué : une boîte de dialogue de confirmation s'affiche, puis le système "
           "crée automatiquement le prochain entretien récurrent.")
add_bullet("Créer un nouvel entretien : via le FAB, formulaire de saisie (véhicule, type, base km ou date).")
add_bullet("Supprimer un entretien.")

add_h3("4.2.6 Carte des véhicules")
add_para(
    "L'onglet Carte (CarteFragment) affiche une carte Google Maps avec un marqueur pour chaque "
    "véhicule de la flotte, positionné selon ses coordonnées GPS stockées en base. "
    "Un clic sur un marqueur affiche une info-bulle avec le nom du véhicule et son statut. "
    "La carte se recentre automatiquement pour englober tous les marqueurs (LatLngBounds)."
)

add_h3("4.2.7 Historique des trajets")
add_para(
    "L'administrateur accède à l'ensemble des trajets réalisés par tous les véhicules, "
    "avec filtres et tri par date, véhicule ou chauffeur."
)

add_h2("4.3 Espace Chauffeur")

add_h3("4.3.1 Accueil")
add_para(
    "L'écran d'accueil (AccueilFragment) accueille le chauffeur par son nom et affiche une "
    "carte récapitulative de son véhicule assigné (marque, modèle, immatriculation, statut). "
    "Un bouton central Démarrer / Arrêter le trajet permet de contrôler le suivi de trajet. "
    "Des raccourcis rapides donnent accès aux fonctions principales."
)

add_h3("4.3.2 Suivi de position")
add_para(
    "L'onglet Position (PositionFragment) affiche la position GPS en temps réel du chauffeur "
    "sur une carte Google Maps. Un bouton Envoyer ma position transmet les coordonnées "
    "actuelles au serveur, mettant à jour les champs lat/lng du véhicule en base de données. "
    "Les permissions de localisation (ACCESS_FINE_LOCATION) sont demandées dynamiquement."
)

add_h3("4.3.3 Mon véhicule")
add_para(
    "Cet écran affiche les informations complètes du véhicule assigné au chauffeur : "
    "photo, marque, modèle, immatriculation, kilométrage, consommation, et les informations "
    "de maintenance (prochaine vidange, prochain contrôle technique)."
)

add_h3("4.3.4 Profil")
add_para(
    "L'onglet Profil (ProfilFragment) présente les informations personnelles du chauffeur "
    "(nom, téléphone, email, permis) et propose un bouton de déconnexion."
)

add_h3("4.3.5 Déclaration d'incident")
add_para(
    "Le chauffeur peut déclarer un incident via DeclarerIncidentActivity. "
    "Le formulaire comprend : sélection du type d'incident (panne mécanique, accident/collision, "
    "pneu crevé, bris de glace, autre), zone de description textuelle, et possibilité "
    "d'attacher plusieurs photos (galerie ou appareil photo). Les photos sont encodées en "
    "Base64 et transmises au serveur dans la liste photos de l'entité Incident. "
    "Les coordonnées GPS actuelles sont automatiquement jointes à la déclaration."
)

add_h3("4.3.6 Historique des trajets personnels")
add_para(
    "Le chauffeur visualise uniquement ses propres trajets : date, heure de départ, "
    "heure d'arrivée, distance parcourue, durée et consommation estimée."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  5. SYSTÈME D'ENTRETIENS ET NOTIFICATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_h1("5. Système d'entretiens et notifications")

add_h2("5.1 Types d'entretiens")
add_para(
    "FleetTracking prend en charge deux modes de déclenchement pour les entretiens préventifs, "
    "contrôlés par le champ estKmBase de l'entité Entretien :"
)
add_bullet("Entretiens kilométriques (estKmBase = true) : déclenchés quand le kilométrage du véhicule "
           "atteint cibleKm. Le type principal est la Vidange (intervalle par défaut : 15 000 km).")
add_bullet("Entretiens temporels (estKmBase = false) : déclenchés à une date cible. "
           "Le type principal est le Contrôle Technique (intervalle par défaut : 365 jours).")

add_h2("5.2 Cycle de vie d'un entretien")
add_para("Lors de la création d'un nouveau véhicule, le backend Spring Boot crée automatiquement :")
add_numbered("Un entretien Vidange avec statut 'aVenir', cibleKm = kilométrage actuel + 15 000.")
add_numbered("Un entretien Contrôle Technique avec statut 'aVenir', cibleDate = date actuelle + 365 jours.")

add_para("Quand un administrateur marque un entretien comme effectué (PUT /api/entretiens/{id}/done) :")
add_numbered("Le statut de l'entretien passe à 'effectue' et dateEffectuee est enregistrée.")
add_numbered("Le backend crée automatiquement le prochain entretien récurrent :")
add_bullet("Pour un entretien km : nouvel entretien avec cibleKm = cibleKm actuel + intervalKm.", level=1)
add_bullet("Pour un entretien date : nouvel entretien avec cibleDate = cibleDate actuelle + intervalJours.", level=1)

add_h2("5.3 Détection des retards et codage couleur")
add_para("La détection de l'urgence s'effectue des deux côtés (Android et backend) :")

add_code_block(
    "// Logique de codage couleur (Android)\n"
    "if (estKmBase) {\n"
    "    double reste = cibleKm - vehiculeKilometrage;\n"
    "    if (reste <= 0)          → badge ROUGE (en retard)\n"
    "    else if (reste <= 500)   → badge ORANGE (imminent)\n"
    "    else                     → badge VERT (OK)\n"
    "} else {\n"
    "    long joursRestants = ChronoUnit.DAYS.between(today, cibleDate);\n"
    "    if (joursRestants <= 0)  → badge ROUGE (en retard)\n"
    "    else if (joursRestants <= 7) → badge ORANGE (imminent)\n"
    "    else                     → badge VERT (OK)\n"
    "}"
)

add_h2("5.4 Système de notifications")
add_para(
    "FleetTracking implémente un système de notifications push locales basé sur WorkManager "
    "pour alerter les utilisateurs des entretiens imminents ou en retard."
)

add_h3("5.4.1 Architecture des notifications")
add_para("Un Worker périodique est déclenché toutes les 24 heures (configurable) :")
add_numbered("Récupération de tous les entretiens avec statut 'aVenir' depuis l'API backend.")
add_numbered("Calcul de l'urgence pour chaque entretien.")
add_numbered("Filtrage selon le profil de l'utilisateur connecté :")
add_bullet("Administrateur : notifications pour tous les véhicules de la flotte.", level=1)
add_bullet("Chauffeur : notifications uniquement pour le véhicule qui lui est assigné.", level=1)
add_numbered("Envoi d'une notification Android locale pour chaque entretien critique.")

add_h3("5.4.2 Canaux de notification")
add_table_styled(
    ["Canal ID", "Nom", "Usage"],
    [
        ["fleet_maintenance", "Entretiens Flotte", "Alertes vidange et contrôle technique"],
        ["fleet_incidents", "Incidents Flotte", "Nouvelles déclarations d'incidents"],
    ]
)

add_h3("5.4.3 Interaction avec les notifications")
add_para(
    "Chaque notification est dotée d'un Intent de redirection : un clic sur la notification "
    "ouvre directement l'écran des entretiens (EntretiensActivity) avec le filtre approprié "
    "(À venir ou En retard). Les notifications utilisent PendingIntent avec le flag "
    "FLAG_IMMUTABLE pour la compatibilité Android 12+."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  6. INTERFACE UTILISATEUR
# ══════════════════════════════════════════════════════════════════════════════
add_h1("6. Interface utilisateur")

add_h2("6.1 Principes de conception")
add_para(
    "L'interface de FleetTracking adopte les guidelines Material Design d'Android pour offrir "
    "une expérience cohérente et intuitive. Les choix de design suivants ont été appliqués :"
)
add_bullet("Palette de couleurs : bleu principal (#1F497D, #2E74B5) pour l'espace admin, palette neutre pour le chauffeur.")
add_bullet("Typography : police Roboto (standard Android) pour les textes, Calibri pour les exports.")
add_bullet("Composants Material : CardView, FloatingActionButton, TabLayout, Spinner, TextInputLayout.")
add_bullet("Navigation : Bottom Navigation Bar pour les deux espaces (Admin et Chauffeur).")

add_h2("6.2 Écran de connexion")
add_para(
    "LoginActivity présente un formulaire épuré : logo de l'application, champs login et mot de passe "
    "(TextInputLayout avec affichage/masquage du mot de passe), RadioGroup pour le choix du rôle "
    "(Admin / Chauffeur), et bouton de connexion. Un ProgressBar s'affiche pendant l'appel API."
)

add_h2("6.3 Espace administrateur")

add_h3("6.3.1 Navigation principale")
add_para(
    "L'AdminActivity héberge une BottomNavigationView avec quatre onglets : "
    "Tableau de bord, Véhicules, Chauffeurs, Carte. "
    "Les fragments correspondants sont chargés dynamiquement dans un NavHostFragment."
)

add_h3("6.3.2 Écran des véhicules (VehiculesFragment)")
add_para(
    "Affiche une liste RecyclerView de cartes véhicules. Chaque carte inclut : "
    "la photo du véhicule (miniature circulaire via Glide), la marque et le modèle, "
    "l'immatriculation, le kilométrage, un badge de statut coloré, et deux boutons "
    "d'action rapide (Incidents, Entretiens). Une SearchView en haut filtre la liste "
    "en temps réel."
)

add_h3("6.3.3 Détail véhicule (VehiculeDetailsActivity)")
add_para(
    "Activité plein écran avec barre d'outils (Toolbar) affichant l'immatriculation. "
    "La photo du véhicule est affichée en grand en haut (avec possibilité de modification). "
    "Les champs éditables sont organisés en groupes logiques dans un ScrollView : "
    "informations générales, état, assignation, maintenance. "
    "Deux boutons en pied de page : Enregistrer et Voir sur la carte."
)

add_h3("6.3.4 Détail chauffeur (ChauffeurDetailsActivity)")
add_para(
    "Similaire à l'activité véhicule : photo de profil en haut, informations personnelles "
    "et professionnelles dans des TextInputLayout Material, spinner d'assignation de véhicule. "
    "En mode création (EXTRA_NEW=true), les champs sont vides et le titre indique 'Nouveau chauffeur'."
)

add_h3("6.3.5 Écran des entretiens")
add_para(
    "TabLayout à deux onglets (À venir / Historique) avec un RecyclerView dans chaque onglet. "
    "Chaque carte d'entretien affiche : type, véhicule, immatriculation, date ou km cible, "
    "badge coloré (vert/orange/rouge), et boutons d'action (Marquer effectué, Supprimer). "
    "Un FAB '+' en bas à droite ouvre le formulaire de création."
)

add_h3("6.3.6 Carte des véhicules (CarteFragment)")
add_para(
    "SupportMapFragment occupant toute la surface de l'onglet. "
    "Les marqueurs sont générés dynamiquement depuis la liste des véhicules. "
    "Un clustering est appliqué pour éviter la superposition des marqueurs. "
    "Un bouton 'Recadrer' repositionne la caméra pour afficher tous les véhicules."
)

add_h2("6.4 Espace chauffeur")

add_h3("6.4.1 Navigation principale")
add_para(
    "La ChauffeurActivity héberge une BottomNavigationView avec quatre onglets : "
    "Accueil, Position, Mon véhicule, Profil."
)

add_h3("6.4.2 Accueil (AccueilFragment)")
add_para(
    "En-tête avec le nom du chauffeur et l'heure. Carte du véhicule assigné avec icône, "
    "modèle et statut. Bouton central de démarrage/arrêt du trajet (vert quand arrêté, "
    "rouge quand en cours, avec chronométre affiché). Grille de raccourcis : "
    "Déclarer un incident, Historique des trajets, Mon profil."
)

add_h3("6.4.3 Position (PositionFragment)")
add_para(
    "Carte Google Maps centrée sur la position actuelle du chauffeur. "
    "Marqueur bleu représentant le chauffeur. Panneau d'informations en bas de l'écran "
    "affichant latitude, longitude, précision et vitesse. Bouton 'Envoyer ma position' "
    "effectuant un PUT sur l'endpoint de mise à jour du véhicule."
)

add_h3("6.4.4 Déclaration d'incident (DeclarerIncidentActivity)")
add_para(
    "Formulaire en plusieurs étapes : Spinner de type d'incident, EditText pour la description, "
    "RecyclerView de miniatures de photos avec bouton d'ajout. Les photos peuvent être "
    "prises avec l'appareil photo ou sélectionnées depuis la galerie. Chaque image est "
    "redimensionnée et compressée (ImageUtils) avant encodage Base64 pour optimiser le "
    "transfert réseau."
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  7. DÉPLOIEMENT ET CONFIGURATION
# ══════════════════════════════════════════════════════════════════════════════
add_h1("7. Déploiement et configuration")

add_h2("7.1 Backend Spring Boot")

add_h3("7.1.1 Configuration application.properties")
add_code_block(
    "# Datasource\n"
    "spring.datasource.url=jdbc:postgresql://localhost:5432/fleettracking\n"
    "spring.datasource.username=fleet_user\n"
    "spring.datasource.password=fleet_password\n"
    "spring.datasource.driver-class-name=org.postgresql.Driver\n\n"
    "# JPA / Hibernate\n"
    "spring.jpa.database-platform=org.hibernate.dialect.PostgreSQLDialect\n"
    "spring.jpa.hibernate.ddl-auto=update\n"
    "spring.jpa.show-sql=false\n\n"
    "# Initialisation des données\n"
    "spring.sql.init.mode=always\n"
    "spring.sql.init.data-locations=classpath:data.sql\n\n"
    "# Port serveur\n"
    "server.port=8080\n\n"
    "# Taille max des requêtes (pour les photos Base64)\n"
    "spring.servlet.multipart.max-file-size=10MB\n"
    "spring.servlet.multipart.max-request-size=50MB"
)

add_h3("7.1.2 Données de démonstration (data.sql)")
add_para(
    "Un fichier data.sql, exécuté au démarrage si la base est vide, insère des données "
    "de démonstration : plusieurs véhicules, chauffeurs, et un utilisateur admin par défaut. "
    "Ce seeding permet de tester l'application sans saisie manuelle."
)

add_h2("7.2 Base de données PostgreSQL via Docker")

add_h3("7.2.1 docker-compose.yml")
add_code_block(
    "version: '3.8'\n"
    "services:\n"
    "  postgres:\n"
    "    image: postgres:15-alpine\n"
    "    container_name: fleettracking-db\n"
    "    environment:\n"
    "      POSTGRES_DB: fleettracking\n"
    "      POSTGRES_USER: fleet_user\n"
    "      POSTGRES_PASSWORD: fleet_password\n"
    "    ports:\n"
    "      - \"5432:5432\"\n"
    "    volumes:\n"
    "      - postgres_data:/var/lib/postgresql/data\n"
    "    restart: unless-stopped\n\n"
    "volumes:\n"
    "  postgres_data:\n"
    "    driver: local"
)

add_h3("7.2.2 Démarrage de la base")
add_code_block(
    "# Démarrage du container PostgreSQL\n"
    "docker-compose up -d\n\n"
    "# Vérification du statut\n"
    "docker-compose ps\n\n"
    "# Connexion à la base pour vérification\n"
    "docker exec -it fleettracking-db psql -U fleet_user -d fleettracking"
)

add_h2("7.3 Application Android")

add_h3("7.3.1 Prérequis")
add_table_styled(
    ["Prérequis", "Version minimale"],
    [
        ["Android Studio", "Hedgehog (2023.1) ou supérieur"],
        ["Android SDK", "API 34 (Android 14) pour le SDK cible"],
        ["minSdkVersion", "API 24 (Android 7.0 Nougat)"],
        ["JDK", "17"],
        ["Gradle", "8.x"],
        ["Google Maps API Key", "Clé valide dans AndroidManifest.xml"],
    ]
)

add_h3("7.3.2 Configuration de l'adresse serveur")
add_para(
    "Avant de compiler l'application, il est nécessaire de configurer l'adresse IP du serveur "
    "backend dans le fichier FleetConfig.java :"
)
add_code_block(
    "// app/src/main/java/com/fleettracking/app/util/FleetConfig.java\n"
    "public class FleetConfig {\n"
    "    // Remplacer par l'adresse IP réelle du serveur Spring Boot\n"
    "    public static final String BASE_URL = \"http://192.168.1.100:8080/api/\";\n\n"
    "    // Timeout en secondes\n"
    "    public static final int CONNECT_TIMEOUT = 30;\n"
    "    public static final int READ_TIMEOUT = 60;\n"
    "}"
)

add_h3("7.3.3 Clé Google Maps")
add_code_block(
    "<!-- AndroidManifest.xml -->\n"
    "<meta-data\n"
    "    android:name=\"com.google.android.geo.API_KEY\"\n"
    "    android:value=\"YOUR_GOOGLE_MAPS_API_KEY\" />"
)

add_h3("7.3.4 Permissions Android")
add_para("Les permissions suivantes sont déclarées dans AndroidManifest.xml :")
add_code_block(
    "<uses-permission android:name=\"android.permission.INTERNET\" />\n"
    "<uses-permission android:name=\"android.permission.ACCESS_FINE_LOCATION\" />\n"
    "<uses-permission android:name=\"android.permission.ACCESS_COARSE_LOCATION\" />\n"
    "<uses-permission android:name=\"android.permission.CAMERA\" />\n"
    "<uses-permission android:name=\"android.permission.READ_EXTERNAL_STORAGE\" />\n"
    "<uses-permission android:name=\"android.permission.POST_NOTIFICATIONS\" />\n"
    "<uses-permission android:name=\"android.permission.RECEIVE_BOOT_COMPLETED\" />"
)

add_h3("7.3.5 Compilation et installation")
add_code_block(
    "# Via Android Studio : Build > Build APK(s)\n"
    "# Ou via ligne de commande :\n"
    "./gradlew assembleDebug\n\n"
    "# Installation sur un appareil connecté :\n"
    "adb install app/build/outputs/apk/debug/app-debug.apk"
)

add_h2("7.4 Procédure de déploiement complète")
add_para("Séquence de démarrage recommandée pour un environnement de développement :")
add_numbered("Cloner le dépôt Git du projet.")
add_numbered("Démarrer PostgreSQL : docker-compose up -d")
add_numbered("Vérifier que le container est healthy : docker-compose ps")
add_numbered("Démarrer le backend Spring Boot : mvn spring-boot:run (ou depuis l'IDE).")
add_numbered("Configurer l'IP du serveur dans FleetConfig.java avec l'IP de la machine hôte.")
add_numbered("Compiler et déployer l'application Android sur l'émulateur ou un appareil physique.")
add_numbered("Vérifier la connectivité réseau (appareil Android et serveur sur le même réseau local).")

add_h2("7.5 Architecture de déploiement réseau")
add_code_block(
    "┌─────────────────────┐         ┌──────────────────────────────┐\n"
    "│  Appareil Android   │         │  Machine de développement    │\n"
    "│  (émulateur ou      │         │                              │\n"
    "│   physique)         │         │  ┌──────────────────────┐   │\n"
    "│                     │  HTTP   │  │  Spring Boot :8080   │   │\n"
    "│  FleetTracking app  ├────────►│  └──────────┬───────────┘   │\n"
    "│                     │         │             │ JDBC           │\n"
    "│  IP: DHCP           │         │  ┌──────────▼───────────┐   │\n"
    "└─────────────────────┘         │  │  PostgreSQL :5432    │   │\n"
    "                                │  │  (Docker container)  │   │\n"
    "    Réseau Wi-Fi / LAN          │  └──────────────────────┘   │\n"
    "                                └──────────────────────────────┘"
)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  8. CONCLUSION / PERSPECTIVES
# ══════════════════════════════════════════════════════════════════════════════
add_h1("8. Conclusion / Perspectives")

add_h2("8.1 Bilan du projet")
add_para(
    "Le projet FleetTracking représente une solution complète et fonctionnelle de gestion "
    "de flotte de véhicules sur mobile. En combinant une application Android moderne "
    "avec un backend Spring Boot robuste et une base de données PostgreSQL conteneurisée, "
    "l'équipe a réussi à livrer un produit couvrant l'ensemble des besoins identifiés :"
)
add_bullet("Authentification sécurisée avec gestion des rôles.")
add_bullet("Suivi GPS en temps réel des véhicules et des chauffeurs.")
add_bullet("Gestion complète du cycle de vie des entretiens préventifs avec récurrence automatique.")
add_bullet("Système de déclaration et de suivi des incidents avec support multi-images.")
add_bullet("Cartographie interactive des positions de la flotte.")
add_bullet("Notifications locales proactives pour les entretiens imminents.")
add_bullet("Enregistrement et consultation des trajets avec calcul automatique des indicateurs.")

add_para(
    "L'architecture choisie (REST client-serveur, séparation claire des couches) garantit "
    "la maintenabilité et l'évolutivité du code. L'utilisation de bibliothèques éprouvées "
    "(Retrofit, WorkManager, Google Maps SDK) assure la stabilité de l'application."
)

add_h2("8.2 Difficultés rencontrées")
add_para("Plusieurs défis techniques ont dû être surmontés lors du développement :")
add_bullet("Gestion des images Base64 : l'encodage/décodage de photos volumineuses a nécessité "
           "une optimisation soignée (compression, redimensionnement) pour éviter les erreurs OOM "
           "(Out Of Memory) sur Android et les limites de taille des requêtes HTTP.")
add_bullet("Synchronisation des données : assurer la cohérence entre le cache local SQLite et "
           "les données serveur lors des reprises de connectivité.")
add_bullet("Permissions Android dynamiques : la gestion des permissions de localisation et de caméra "
           "sur les versions récentes d'Android (12+) a requis une attention particulière.")
add_bullet("Calcul de la récurrence des entretiens : la logique de création automatique des "
           "prochains entretiens a nécessité une validation rigoureuse des cas limites.")
add_bullet("Configuration réseau : l'adresse IP du serveur devant être configurée manuellement, "
           "le déploiement dans des environnements variés peut être source d'erreurs.")

add_h2("8.3 Perspectives d'amélioration")
add_para(
    "Plusieurs axes d'amélioration ont été identifiés pour faire évoluer FleetTracking "
    "vers une solution de niveau production :"
)

add_h3("8.3.1 Améliorations techniques")
add_bullet("Authentification JWT complète : implémenter un système de refresh token avec expiration "
           "et rotation automatique des tokens de sécurité.")
add_bullet("Discovery de services : remplacer l'IP statique par un mécanisme de découverte automatique "
           "du serveur sur le réseau local (mDNS / Bonjour), ou utiliser un nom de domaine.")
add_bullet("Chiffrement HTTPS : déployer un certificat TLS pour sécuriser les communications en production.")
add_bullet("Mode hors-ligne avancé : synchronisation bidirectionnelle avec résolution de conflits "
           "pour une utilisation sans connectivité prolongée.")
add_bullet("Architecture MVVM : migrer vers le pattern ViewModel + LiveData / StateFlow pour "
           "une meilleure séparation des responsabilités et une meilleure testabilité.")
add_bullet("Tests automatisés : développer une suite de tests unitaires (JUnit) et d'intégration "
           "(Espresso pour Android, MockMvc pour Spring Boot).")

add_h3("8.3.2 Nouvelles fonctionnalités")
add_bullet("Tableau de bord analytique : graphiques de consommation, coûts d'entretien, "
           "statistiques de trajets avec export PDF/Excel.")
add_bullet("Géofencing : alertes lorsqu'un véhicule sort d'une zone géographique définie.")
add_bullet("Notifications push FCM : remplacer les notifications locales WorkManager par "
           "Firebase Cloud Messaging pour des alertes en temps réel depuis le serveur.")
add_bullet("Rapport d'incident automatisé : génération de PDF de déclaration d'incident.")
add_bullet("Application web d'administration : interface web React/Angular pour les responsables "
           "n'ayant pas accès à l'application mobile.")
add_bullet("Intégration comptabilité : export des données de trajet et de consommation "
           "vers des systèmes comptables (format CSV, API SAP/Sage).")
add_bullet("Intelligence artificielle : prédiction des pannes basée sur l'historique de maintenance "
           "et les données de capteurs OBD-II.")

add_h3("8.3.3 Déploiement en production")
add_bullet("Conteneurisation complète : Docker Compose incluant le backend Spring Boot "
           "(Dockerfile Maven multi-stage build).")
add_bullet("Cloud deployment : déploiement sur AWS, Azure ou GCP avec scaling horizontal.")
add_bullet("CI/CD : pipeline GitHub Actions pour build, tests et déploiement automatisés.")
add_bullet("Monitoring : intégration Prometheus/Grafana pour le suivi des métriques serveur.")

add_h2("8.4 Conclusion générale")
add_para(
    "FleetTracking démontre la viabilité technique et fonctionnelle d'une solution mobile "
    "de gestion de flotte développée dans un contexte universitaire. Les technologies Android, "
    "Spring Boot et PostgreSQL, combinées selon une architecture REST claire, offrent une "
    "base solide sur laquelle des fonctionnalités avancées pourront être construites. "
    "Ce projet a permis à l'équipe d'acquérir des compétences approfondies en développement "
    "mobile Android, en conception d'API REST, en gestion de base de données relationnelle, "
    "et en containerisation Docker — compétences directement transférables dans un contexte "
    "professionnel."
)

add_para(
    "La modularité de l'architecture et la qualité du code produit constituent un socle "
    "technologique sain pour une éventuelle mise en production ou pour une continuation "
    "du projet dans le cadre d'autres modules universitaires."
)

# ── Save ──────────────────────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Rapport généré avec succès : {OUTPUT_PATH}")
