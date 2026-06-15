"""
Generate 3 diagrams (Use Case, Class, Sequence) as PNG images using matplotlib,
then insert them into the existing rapport_projet.docx at the appropriate sections.
"""

import os
import sys
import tempfile
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Ellipse, FancyArrowPatch
import matplotlib.patheffects as pe
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOCX_PATH = r'C:\Users\danad\Documents\Claude\Projects\projetmobile\docs\rapport_projet.docx'

# ─────────────────────────────────────────────
# 1. USE CASE DIAGRAM
# ─────────────────────────────────────────────

def draw_stick_figure(ax, x, y, label, color='#2c3e50'):
    """Draw a simple stick figure actor."""
    # Head
    head = plt.Circle((x, y + 0.55), 0.18, color=color, zorder=5)
    ax.add_patch(head)
    # Body
    ax.plot([x, x], [y + 0.37, y - 0.15], color=color, lw=2, zorder=5)
    # Arms
    ax.plot([x - 0.28, x + 0.28], [y + 0.15, y + 0.15], color=color, lw=2, zorder=5)
    # Legs
    ax.plot([x, x - 0.22], [y - 0.15, y - 0.55], color=color, lw=2, zorder=5)
    ax.plot([x, x + 0.22], [y - 0.15, y - 0.55], color=color, lw=2, zorder=5)
    # Label
    ax.text(x, y - 0.82, label, ha='center', va='top', fontsize=11,
            fontweight='bold', color=color, zorder=5)


def draw_use_case(ax, x, y, text, width=2.4, height=0.55,
                  facecolor='#eaf4fb', edgecolor='#2980b9', fontsize=9.5):
    """Draw an oval use case."""
    ellipse = Ellipse((x, y), width=width, height=height,
                      facecolor=facecolor, edgecolor=edgecolor, lw=1.5, zorder=4)
    ax.add_patch(ellipse)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            color='#1a252f', zorder=5, wrap=True,
            multialignment='center')


def generate_use_case_diagram(path):
    fig, ax = plt.subplots(figsize=(16, 11))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    # System boundary
    rect = FancyBboxPatch((2.6, 0.5), 10.8, 9.8,
                          boxstyle="round,pad=0.15",
                          facecolor='#f0f8ff', edgecolor='#2c3e50',
                          lw=2.5, zorder=1)
    ax.add_patch(rect)
    ax.text(8.0, 10.55, 'FleetTracking', ha='center', va='center',
            fontsize=14, fontweight='bold', color='#2c3e50',
            bbox=dict(facecolor='white', edgecolor='none', pad=2))

    # ── Admin (left) ──
    admin_x = 1.1
    admin_y = 5.5
    draw_stick_figure(ax, admin_x, admin_y, 'Admin', color='#1a5276')

    admin_cases = [
        (5.3, 9.2, 'Gérer les\nvéhicules'),
        (5.3, 7.9, 'Gérer les\nchauffeurs'),
        (5.3, 6.6, 'Consulter\nla carte'),
        (5.3, 5.3, 'Gérer les\nincidents'),
        (5.3, 4.0, 'Gérer les\nentretiens'),
        (5.3, 2.7, 'Créer un\nentretien'),
        (5.3, 1.4, "Consulter\nl'historique"),
    ]
    for (cx, cy, label) in admin_cases:
        draw_use_case(ax, cx, cy, label)
        ax.annotate('', xy=(cx - 1.22, cy), xytext=(admin_x + 0.3, admin_y),
                    arrowprops=dict(arrowstyle='-', color='#566573', lw=1.3),
                    zorder=3)

    # ── Chauffeur (right) ──
    chauffeur_x = 14.9
    chauffeur_y = 5.5
    draw_stick_figure(ax, chauffeur_x, chauffeur_y, 'Chauffeur', color='#117a65')

    chauffeur_cases = [
        (10.7, 8.5, 'Démarrer/Arrêter\nun trajet'),
        (10.7, 7.1, 'Envoyer\nsa position'),
        (10.7, 5.7, 'Déclarer\nun incident'),
        (10.7, 4.3, 'Consulter\nson véhicule'),
        (10.7, 2.9, 'Voir ses\nentretiens'),
    ]
    for (cx, cy, label) in chauffeur_cases:
        draw_use_case(ax, cx, cy, label, facecolor='#e9f7ef', edgecolor='#1e8449')
        ax.annotate('', xy=(cx + 1.22, cy), xytext=(chauffeur_x - 0.3, chauffeur_y),
                    arrowprops=dict(arrowstyle='-', color='#566573', lw=1.3),
                    zorder=3)

    ax.set_title("Diagramme des cas d'utilisation — FleetTracking",
                 fontsize=15, fontweight='bold', pad=14, color='#1a252f')
    plt.tight_layout(pad=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  Use case diagram saved: {path}")


# ─────────────────────────────────────────────
# 2. CLASS DIAGRAM
# ─────────────────────────────────────────────

def draw_class_box(ax, x, y, name, fields, width=3.0, row_h=0.32,
                   header_color='#2471a3', body_color='#eaf4fb'):
    """Draw a UML-style class box."""
    n = len(fields)
    total_h = row_h + n * row_h + 0.1  # header + fields + padding

    # Header
    header = FancyBboxPatch((x - width / 2, y - row_h / 2), width, row_h,
                            boxstyle="square,pad=0", facecolor=header_color,
                            edgecolor='#1a5276', lw=1.8, zorder=4)
    ax.add_patch(header)
    ax.text(x, y, name, ha='center', va='center', fontsize=10,
            fontweight='bold', color='white', zorder=5)

    # Fields
    body = FancyBboxPatch((x - width / 2, y - row_h / 2 - n * row_h), width, n * row_h,
                          boxstyle="square,pad=0", facecolor=body_color,
                          edgecolor='#1a5276', lw=1.8, zorder=4)
    ax.add_patch(body)
    for i, field in enumerate(fields):
        fy = y - row_h / 2 - (i + 0.5) * row_h
        ax.text(x - width / 2 + 0.12, fy, f'  {field}', ha='left', va='center',
                fontsize=8.2, color='#1a252f', zorder=5, family='monospace')

    bottom_y = y - row_h / 2 - n * row_h
    return bottom_y, total_h  # bottom edge, total height


def connect_boxes(ax, x1, y1_top, x2, y2_top, label1='', label2='',
                  color='#2c3e50', style='-'):
    """Draw a line between two class box center-tops."""
    ax.annotate('', xy=(x2, y2_top), xytext=(x1, y1_top),
                arrowprops=dict(arrowstyle='-', color=color, lw=1.5,
                                connectionstyle='arc3,rad=0.0'),
                zorder=3)
    mid_x = (x1 + x2) / 2
    mid_y = (y1_top + y2_top) / 2
    if label1:
        ax.text(x1 + (x2 - x1) * 0.12, y1_top + (y2_top - y1_top) * 0.12,
                label1, fontsize=8, color='#922b21', fontweight='bold', zorder=6,
                ha='center', va='bottom')
    if label2:
        ax.text(x1 + (x2 - x1) * 0.88, y1_top + (y2_top - y1_top) * 0.88,
                label2, fontsize=8, color='#922b21', fontweight='bold', zorder=6,
                ha='center', va='bottom')


def generate_class_diagram(path):
    fig, ax = plt.subplots(figsize=(18, 11))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    row_h = 0.34
    w = 3.3

    # ── Define entities (cx, cy=top-of-header) ──
    entities = {
        'Vehicule': {
            'x': 9.0, 'y': 9.5,
            'fields': ['id: Long', 'marque: String', 'modele: String',
                       'immatriculation: String', 'kilometrage: Double',
                       'statut: String', 'conducteurId: Long'],
            'color': '#1a5276',
        },
        'Chauffeur': {
            'x': 3.0, 'y': 9.5,
            'fields': ['id: Long', 'nom: String', 'telephone: String',
                       'email: String', 'statut: String'],
            'color': '#117a65',
        },
        'Entretien': {
            'x': 14.5, 'y': 6.8,
            'fields': ['id: Long', 'vehiculeId: Long', 'type: String',
                       'estKmBase: Boolean', 'cibleKm: Double',
                       'cibleDate: Date', 'statut: String'],
            'color': '#6e2fa0',
        },
        'Incident': {
            'x': 9.0, 'y': 4.2,
            'fields': ['id: Long', 'vehiculeId: Long', 'type: String',
                       'description: String', 'statut: String'],
            'color': '#922b21',
        },
        'Trajet': {
            'x': 3.0, 'y': 4.8,
            'fields': ['id: Long', 'vehiculeId: Long', 'chauffeurId: Long',
                       'distanceKm: Double', 'duree: Long'],
            'color': '#784212',
        },
    }

    boxes = {}  # name -> (center_x, top_y, bottom_y)
    for name, info in entities.items():
        cx, cy = info['x'], info['y']
        bottom, _ = draw_class_box(ax, cx, cy, name, info['fields'],
                                   width=w, row_h=row_h,
                                   header_color=info['color'])
        boxes[name] = (cx, cy, bottom)

    def mid_y(name):
        cx, top, bot = boxes[name]
        return (top + bot) / 2

    # ── Relationships ──
    # Vehicule 1 — 0..* Entretien
    vx, vt, vb = boxes['Vehicule']
    ex, et, eb = boxes['Entretien']
    ax.annotate('', xy=(ex - w / 2, et), xytext=(vx + w / 2, vt),
                arrowprops=dict(arrowstyle='-', color='#2c3e50', lw=1.8), zorder=3)
    ax.text(vx + w / 2 + 0.1, vt + 0.05, '1', fontsize=9, color='#922b21', fontweight='bold')
    ax.text(ex - w / 2 - 0.18, et + 0.05, '0..*', fontsize=9, color='#922b21', fontweight='bold')
    ax.text((vx + w / 2 + ex - w / 2) / 2, (vt + et) / 2 + 0.15,
            'vehiculeId', fontsize=8, color='#555', ha='center', style='italic')

    # Vehicule 1 — 0..* Incident
    ix, it, ib = boxes['Incident']
    ax.annotate('', xy=(ix + w / 2, it), xytext=(vx + w / 2, vb),
                arrowprops=dict(arrowstyle='-', color='#2c3e50', lw=1.8), zorder=3)
    ax.text(vx + w / 2 + 0.1, vb + 0.05, '1', fontsize=9, color='#922b21', fontweight='bold')
    ax.text(ix + w / 2 + 0.08, it - 0.15, '0..*', fontsize=9, color='#922b21', fontweight='bold')
    ax.text((vx + w / 2 + ix + w / 2) / 2 + 0.5, (vb + it) / 2,
            'vehiculeId', fontsize=8, color='#555', ha='center', style='italic')

    # Vehicule 1 — 0..* Trajet
    tx, tt, tb = boxes['Trajet']
    ax.annotate('', xy=(tx + w / 2, tt), xytext=(vx - w / 2, vb + (vt - vb) * 0.3),
                arrowprops=dict(arrowstyle='-', color='#2c3e50', lw=1.8), zorder=3)
    ax.text(vx - w / 2 - 0.18, vb + (vt - vb) * 0.32, '1', fontsize=9,
            color='#922b21', fontweight='bold')
    ax.text(tx + w / 2 + 0.08, tt + 0.05, '0..*', fontsize=9, color='#922b21', fontweight='bold')
    ax.text((vx - w / 2 + tx + w / 2) / 2, (vb + tt) / 2 + 0.18,
            'vehiculeId', fontsize=8, color='#555', ha='center', style='italic')

    # Chauffeur 0..1 — 0..1 Vehicule
    cx2, ct, cb = boxes['Chauffeur']
    ax.annotate('', xy=(vx - w / 2, vt), xytext=(cx2 + w / 2, ct),
                arrowprops=dict(arrowstyle='-', color='#2c3e50', lw=1.8), zorder=3)
    ax.text(cx2 + w / 2 + 0.08, ct + 0.05, '0..1', fontsize=9, color='#922b21', fontweight='bold')
    ax.text(vx - w / 2 - 0.22, vt + 0.05, '0..1', fontsize=9, color='#922b21', fontweight='bold')
    ax.text((cx2 + w / 2 + vx - w / 2) / 2, (ct + vt) / 2 + 0.15,
            'conducteurId', fontsize=8, color='#555', ha='center', style='italic')

    # Chauffeur 1 — 0..* Trajet
    ax.annotate('', xy=(tx, tt), xytext=(cx2, cb),
                arrowprops=dict(arrowstyle='-', color='#2c3e50', lw=1.8), zorder=3)
    ax.text(cx2 - 0.12, cb - 0.12, '1', fontsize=9, color='#922b21', fontweight='bold')
    ax.text(tx - 0.12, tt + 0.08, '0..*', fontsize=9, color='#922b21', fontweight='bold')
    ax.text((cx2 + tx) / 2 - 0.4, (cb + tt) / 2,
            'chauffeurId', fontsize=8, color='#555', ha='center', style='italic')

    ax.set_title('Diagramme de classes — FleetTracking',
                 fontsize=15, fontweight='bold', pad=14, color='#1a252f')
    plt.tight_layout(pad=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  Class diagram saved: {path}")


# ─────────────────────────────────────────────
# 3. SEQUENCE DIAGRAM
# ─────────────────────────────────────────────

def generate_sequence_diagram(path):
    fig, ax = plt.subplots(figsize=(16, 11))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    # Participants
    participants = [
        ('Admin/Chauffeur', 1.5,  '#1a5276'),
        ('Android App',     5.2,  '#117a65'),
        ('Backend\n(Spring Boot)', 9.5,  '#6e2fa0'),
        ('PostgreSQL',      13.8, '#922b21'),
    ]

    header_y = 10.3
    lifeline_top = 10.05
    lifeline_bottom = 0.3

    # Draw participant boxes
    for name, x, color in participants:
        box = FancyBboxPatch((x - 1.05, header_y - 0.28), 2.1, 0.58,
                             boxstyle="round,pad=0.07",
                             facecolor=color, edgecolor='#1a252f', lw=1.5, zorder=5)
        ax.add_patch(box)
        ax.text(x, header_y, name, ha='center', va='center', fontsize=9.5,
                fontweight='bold', color='white', zorder=6, multialignment='center')

        # Lifeline (dashed vertical)
        ax.plot([x, x], [lifeline_top - 0.3, lifeline_bottom],
                color='#717d7e', lw=1.3, linestyle='--', zorder=2, dashes=(5, 4))

    # Activation boxes
    act_boxes = [
        # (participant_index, y_top, y_bottom)
        (1, 9.6, 0.4),    # Android App active most of the time
        (2, 7.9, 5.6),    # Backend active during server processing
        (3, 7.6, 5.9),    # PostgreSQL active during DB ops
    ]
    act_color = '#d5d8dc'
    for pidx, yt, yb in act_boxes:
        px = participants[pidx][1]
        box = FancyBboxPatch((px - 0.14, yb), 0.28, yt - yb,
                             boxstyle="square,pad=0",
                             facecolor=act_color, edgecolor='#717d7e', lw=1.2, zorder=3)
        ax.add_patch(box)

    def arrow(x1, x2, y, label, color='#2c3e50', is_return=False, self_call=False):
        style = '-->' if is_return else '->'
        ax.annotate('',
                    xy=(x2, y), xytext=(x1, y),
                    arrowprops=dict(
                        arrowstyle='->' if not is_return else '->',
                        color=color, lw=1.6,
                        linestyle='dashed' if is_return else 'solid',
                    ), zorder=4)
        mid_x = (x1 + x2) / 2
        offset = 0.13
        ax.text(mid_x, y + offset, label, ha='center', va='bottom',
                fontsize=8.5, color='#1a252f', zorder=5,
                bbox=dict(facecolor='white', edgecolor='none', alpha=0.7, pad=1))

    pxs = [p[1] for p in participants]

    # Step numbers Y positions (top to bottom)
    steps_y = [9.35, 8.85, 8.35, 7.85, 7.35, 6.85, 6.35, 5.85, 5.2, 4.6]

    arrows = [
        # (from_idx, to_idx, label, is_return)
        (0, 1, 'Appuie sur "Marquer effectué"',          False),
        (1, 1, 'Affiche dialog de confirmation',          False),  # self-call
        (0, 1, 'Confirme',                                False),
        (1, 2, 'PUT /api/entretiens/{id}/done',           False),
        (2, 3, "UPDATE entretien SET statut='effectue'",  False),
        (2, 3, 'INSERT INTO entretien (nouveau, récurrent)', False),
        (3, 2, 'OK',                                      True),
        (2, 1, '200 OK + Entretien mis à jour',           True),
        (1, 1, 'Recharge la liste',                       False),  # self
        (1, 0, 'Affiche liste mise à jour',               True),
    ]

    for i, (fi, ti, label, ret) in enumerate(arrows):
        y = steps_y[i]
        x1 = pxs[fi]
        x2 = pxs[ti]

        # Step number circle
        ax.text(0.18, y, str(i + 1), ha='center', va='center', fontsize=7.5,
                color='white', zorder=7,
                bbox=dict(boxstyle='circle,pad=0.2', facecolor='#2c3e50',
                          edgecolor='none'))

        if fi == ti:
            # Self-call loop
            loop_x = x1 + 0.14
            loop_w = 0.7
            loop_h = 0.35
            rect = mpatches.FancyArrowPatch(
                (loop_x, y), (loop_x + loop_w, y - loop_h),
                arrowstyle='->', color='#2c3e50', lw=1.5,
                connectionstyle='arc3,rad=-0.5', zorder=4,
            )
            ax.add_patch(rect)
            ax.text(loop_x + loop_w / 2 + 0.45, y - loop_h / 2, label,
                    ha='left', va='center', fontsize=8.2, color='#1a252f', zorder=5)
        else:
            # Horizontal arrow
            linestyle = 'dashed' if ret else 'solid'
            color = '#717d7e' if ret else '#1a252f'
            ax.annotate('', xy=(x2, y), xytext=(x1, y),
                        arrowprops=dict(arrowstyle='->', color=color, lw=1.5,
                                        linestyle=linestyle), zorder=4)
            mid_x = (x1 + x2) / 2
            ax.text(mid_x, y + 0.12, label, ha='center', va='bottom',
                    fontsize=8.3, color=color, zorder=5,
                    bbox=dict(facecolor='white', edgecolor='none', alpha=0.75, pad=1))

    ax.set_title("Diagramme de séquence — Marquer un entretien comme effectué",
                 fontsize=14, fontweight='bold', pad=14, color='#1a252f')
    plt.tight_layout(pad=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  Sequence diagram saved: {path}")


# ─────────────────────────────────────────────
# DOCX INSERTION HELPERS
# ─────────────────────────────────────────────

def find_paragraph_index(doc, heading_text_contains):
    """Return the index of the first paragraph whose text contains the given string."""
    for i, p in enumerate(doc.paragraphs):
        if heading_text_contains.lower() in p.text.lower():
            return i
    return None


def insert_image_after_paragraph(doc, para_index, image_path, width_inches=6.0):
    """
    Insert an image paragraph immediately after the paragraph at para_index.
    Uses python-docx's internal XML manipulation to place it at the right position.
    """
    # We need to insert a new paragraph containing the image right after para_index.
    # python-docx doesn't have a direct insert_after, so we manipulate the XML.
    target_para = doc.paragraphs[para_index]

    # Create a new paragraph element with the picture
    # First add it at the end then move it
    new_para = doc.add_paragraph()
    new_para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    # Move the new paragraph to be right after target_para in the XML
    target_para._element.addnext(new_para._element)

    # Also add a small caption paragraph
    caption_para = doc.add_paragraph()
    caption_para.alignment = 1
    caption_run = caption_para.add_run()
    caption_run.font.size = None  # default
    caption_run.font.italic = True
    caption_run.font.color.rgb = None

    # Move caption right after the image paragraph
    new_para._element.addnext(caption_para._element)

    return new_para


def insert_image_after_heading(doc, heading_contains, image_path, width_inches=6.0):
    """Find a heading and insert an image paragraph right after it."""
    idx = find_paragraph_index(doc, heading_contains)
    if idx is None:
        print(f"  WARNING: heading not found: '{heading_contains}'")
        return False
    print(f"  Inserting image after paragraph {idx}: '{doc.paragraphs[idx].text[:60]}'")
    insert_image_after_paragraph(doc, idx, image_path, width_inches)
    return True


# ─────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────

def main():
    tmp_dir = tempfile.mkdtemp()
    uc_path  = os.path.join(tmp_dir, 'use_case.png')
    cls_path = os.path.join(tmp_dir, 'class.png')
    seq_path = os.path.join(tmp_dir, 'sequence.png')

    print("Generating diagrams...")
    generate_use_case_diagram(uc_path)
    generate_class_diagram(cls_path)
    generate_sequence_diagram(seq_path)

    print(f"\nOpening document: {DOCX_PATH}")
    doc = Document(DOCX_PATH)

    # Insert Use Case diagram → after "4. Fonctionnalités principales"
    print("\nInserting Use Case diagram...")
    insert_image_after_heading(doc, '4. Fonctionnalit', uc_path, width_inches=6.0)

    # Insert Class diagram → after "3. Modèle de données"
    print("Inserting Class diagram...")
    insert_image_after_heading(doc, '3. Mod', cls_path, width_inches=6.0)

    # Insert Sequence diagram → after "5. Système d'entretiens et notifications"
    print("Inserting Sequence diagram...")
    insert_image_after_heading(doc, "5. Syst", seq_path, width_inches=6.0)

    print(f"\nSaving document to: {DOCX_PATH}")
    doc.save(DOCX_PATH)
    print("Done!")

    # Clean up temp files
    for f in [uc_path, cls_path, seq_path]:
        try:
            os.remove(f)
        except Exception:
            pass
    try:
        os.rmdir(tmp_dir)
    except Exception:
        pass
    print("Temp files cleaned up.")


if __name__ == '__main__':
    main()
