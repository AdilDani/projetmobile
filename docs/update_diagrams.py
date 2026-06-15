"""
update_diagrams.py
------------------
1. Opens the existing rapport_projet.docx
2. Removes paragraphs that contain InlineShapes (old embedded images)
3. Generates 4 new diagrams as PNGs using matplotlib
4. Inserts them in the correct sections
5. Saves the docx
"""

import os
import sys
import tempfile
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Ellipse, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX_PATH = r'C:\Users\danad\Documents\Claude\Projects\projetmobile\docs\rapport_projet.docx'

# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------

def paragraph_has_image(para):
    """Return True if the paragraph XML contains an inline image."""
    xml = para._element.xml
    return 'pic:pic' in xml or 'a:blip' in xml or 'w:drawing' in xml


def remove_image_paragraphs(doc):
    """Remove all paragraphs that contain embedded images from the document body."""
    removed = 0
    body = doc.element.body
    paras_to_remove = [p for p in doc.paragraphs if paragraph_has_image(p)]
    for p in paras_to_remove:
        p._element.getparent().remove(p._element)
        removed += 1
    print(f"  Removed {removed} image paragraph(s).")
    return removed


def find_paragraph_index(doc, *keywords, style_contains=None):
    """
    Return the index of the first paragraph whose text contains ALL keywords
    (case-insensitive). Optionally filter by style name.
    """
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.lower()
        if all(kw.lower() in txt for kw in keywords):
            if style_contains is None or style_contains.lower() in p.style.name.lower():
                return i
    return None


def insert_image_after_paragraph(doc, para_index, image_path, width_inches=6.5,
                                  caption=None):
    """
    Insert an image paragraph immediately after the paragraph at para_index.
    Uses XML manipulation to place it at the right position.
    """
    target_para = doc.paragraphs[para_index]

    # Image paragraph
    new_para = doc.add_paragraph()
    new_para.alignment = 1  # CENTER
    run = new_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    target_para._element.addnext(new_para._element)

    # Optional caption paragraph right after image
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = 1
        cap_run = cap_para.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        new_para._element.addnext(cap_para._element)

    return new_para


# ---------------------------------------------------------------------------
# DIAGRAM 1 — USE CASE
# ---------------------------------------------------------------------------

def draw_stick_figure(ax, x, y, label, color='#2c3e50'):
    head = plt.Circle((x, y + 0.55), 0.20, color=color, zorder=5)
    ax.add_patch(head)
    ax.plot([x, x],          [y + 0.35, y - 0.15], color=color, lw=2.5, zorder=5)
    ax.plot([x - 0.32, x + 0.32], [y + 0.12, y + 0.12], color=color, lw=2.5, zorder=5)
    ax.plot([x, x - 0.26],  [y - 0.15, y - 0.60], color=color, lw=2.5, zorder=5)
    ax.plot([x, x + 0.26],  [y - 0.15, y - 0.60], color=color, lw=2.5, zorder=5)
    ax.text(x, y - 0.90, label, ha='center', va='top', fontsize=11,
            fontweight='bold', color=color, zorder=5)


def draw_use_case_oval(ax, x, y, text, width=2.6, height=0.60,
                       facecolor='#eaf4fb', edgecolor='#2980b9'):
    ellipse = Ellipse((x, y), width=width, height=height,
                      facecolor=facecolor, edgecolor=edgecolor, lw=1.8, zorder=4)
    ax.add_patch(ellipse)
    ax.text(x, y, text, ha='center', va='center', fontsize=9,
            color='#1a252f', zorder=5, multialignment='center')


def generate_use_case_diagram(path):
    fig, ax = plt.subplots(figsize=(16, 11))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    # System boundary
    rect = FancyBboxPatch((2.8, 0.5), 10.4, 9.8,
                          boxstyle="round,pad=0.15",
                          facecolor='#f0f8ff', edgecolor='#2c3e50',
                          lw=2.5, zorder=1)
    ax.add_patch(rect)
    ax.text(8.0, 10.55, 'Système FleetTracking',
            ha='center', va='center', fontsize=14, fontweight='bold', color='#2c3e50',
            bbox=dict(facecolor='white', edgecolor='none', pad=3))

    # ── Admin (left actor) ──
    admin_x, admin_y = 1.1, 5.5
    draw_stick_figure(ax, admin_x, admin_y, 'Admin', color='#1a5276')

    admin_cases = [
        (5.3, 9.2, 'Gérer les\nvéhicules'),
        (5.3, 8.0, 'Gérer les\nchauffeurs'),
        (5.3, 6.8, 'Consulter la carte\nen temps réel'),
        (5.3, 5.5, 'Gérer les\nincidents'),
        (5.3, 4.2, 'Gérer les\nentretiens'),
        (5.3, 3.0, 'Créer un\nentretien'),
        (5.3, 1.8, "Consulter l'historique\ndes trajets"),
        (5.3, 0.9, 'Marquer un entretien\neffectué'),
    ]
    for (cx, cy, label) in admin_cases:
        draw_use_case_oval(ax, cx, cy, label,
                           facecolor='#eaf4fb', edgecolor='#2980b9')
        ax.annotate('', xy=(cx - 1.32, cy), xytext=(admin_x + 0.32, admin_y),
                    arrowprops=dict(arrowstyle='-', color='#566573', lw=1.3), zorder=3)

    # ── Chauffeur (right actor) ──
    chauffeur_x, chauffeur_y = 14.9, 5.5
    draw_stick_figure(ax, chauffeur_x, chauffeur_y, 'Chauffeur', color='#117a65')

    chauffeur_cases = [
        (10.7, 8.7, 'Démarrer\nun trajet'),
        (10.7, 7.5, 'Arrêter\nun trajet'),
        (10.7, 6.2, 'Envoyer\nsa position'),
        (10.7, 4.9, 'Déclarer\nun incident'),
        (10.7, 3.6, 'Consulter\nson véhicule'),
        (10.7, 2.3, 'Consulter\nses entretiens'),
    ]
    for (cx, cy, label) in chauffeur_cases:
        draw_use_case_oval(ax, cx, cy, label,
                           facecolor='#e9f7ef', edgecolor='#1e8449')
        ax.annotate('', xy=(cx + 1.32, cy), xytext=(chauffeur_x - 0.32, chauffeur_y),
                    arrowprops=dict(arrowstyle='-', color='#566573', lw=1.3), zorder=3)

    ax.set_title("Diagramme des cas d'utilisation — FleetTracking",
                 fontsize=15, fontweight='bold', pad=14, color='#1a252f')
    plt.tight_layout(pad=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  [1] Use Case diagram saved: {path}")


# ---------------------------------------------------------------------------
# DIAGRAM 2 — CLASS DIAGRAM
# ---------------------------------------------------------------------------

def draw_class_box(ax, cx, cy, name, fields, width=3.2, row_h=0.28,
                   header_color='#2471a3', body_color='#eaf4fb'):
    """Draw UML class box. Returns (top_y, bottom_y)."""
    n = len(fields)
    total_fields_h = n * row_h
    header_h = 0.36

    # Header rect
    hx = cx - width / 2
    hy = cy - header_h / 2
    hrect = FancyBboxPatch((hx, hy), width, header_h,
                           boxstyle="square,pad=0", facecolor=header_color,
                           edgecolor='#1a252f', lw=1.8, zorder=4)
    ax.add_patch(hrect)
    ax.text(cx, cy, name, ha='center', va='center', fontsize=9.5,
            fontweight='bold', color='white', zorder=5)

    # Body rect
    body_top = cy - header_h / 2
    body_bottom = body_top - total_fields_h
    brect = FancyBboxPatch((hx, body_bottom), width, total_fields_h,
                           boxstyle="square,pad=0", facecolor=body_color,
                           edgecolor='#1a252f', lw=1.8, zorder=4)
    ax.add_patch(brect)
    for i, field in enumerate(fields):
        fy = body_top - (i + 0.5) * row_h
        ax.text(hx + 0.10, fy, field, ha='left', va='center',
                fontsize=7.2, color='#1a252f', zorder=5, family='monospace')

    top_y = cy + header_h / 2
    bottom_y = body_bottom
    return top_y, bottom_y


def generate_class_diagram(path):
    fig, ax = plt.subplots(figsize=(20, 13))
    ax.set_xlim(0, 20)
    ax.set_ylim(0, 13)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    row_h = 0.265
    w = 3.5

    # ── Entity definitions ──
    entities = {
        'Vehicule': {
            'cx': 9.5, 'cy': 11.8,
            'color': '#1a5276',
            'fields': [
                'id: String',
                'marque: String',
                'modele: String',
                'immatriculation: String',
                'annee: int',
                'kilometrage: int',
                'statut: String',
                'carburantPct: int',
                'consommation: double',
                'vidangeCibleKm: int',
                'controleTechniqueDate: String',
                'conducteurId: String',
                'lat: double',
                'lng: double',
                'vitesse: int',
                'photo: String (base64)',
            ],
        },
        'Chauffeur': {
            'cx': 3.0, 'cy': 11.5,
            'color': '#117a65',
            'fields': [
                'id: String',
                'nom: String',
                'telephone: String',
                'email: String',
                'permis: String',
                'login: String',
                'motDePasse: String',
                'statut: String',
                'photo: String (base64)',
            ],
        },
        'AdminUser': {
            'cx': 16.5, 'cy': 11.5,
            'color': '#6e2fa0',
            'fields': [
                'id: String',
                'login: String',
                'motDePasse: String',
            ],
        },
        'Entretien': {
            'cx': 3.0, 'cy': 5.5,
            'color': '#784212',
            'fields': [
                'id: String',
                'vehiculeId: String',
                'vehiculeNom: String',
                'immatriculation: String',
                'type: String',
                'estKmBase: boolean',
                'cibleKm: int',
                'cibleDate: String',
                'intervalKm: int',
                'intervalJours: int',
                'statut: String',
                'dateEffectuee: String',
                'retardJours: int',
            ],
        },
        'Incident': {
            'cx': 9.5, 'cy': 5.5,
            'color': '#922b21',
            'fields': [
                'id: String',
                'vehiculeId: String',
                'vehiculeNom: String',
                'type: String',
                'description: String',
                'date: String',
                'statut: String',
                'photos: String (JSON)',
                'lat: double',
                'lng: double',
            ],
        },
        'Trajet': {
            'cx': 16.5, 'cy': 6.0,
            'color': '#1a5276',
            'fields': [
                'id: String',
                'vehiculeId: String',
                'vehiculeNom: String',
                'chauffeurId: String',
                'date: String',
                'heureDepart: String',
                'heureArrivee: String',
                'distanceKm: int',
                'duree: String',
                'vitesseMoyenne: int',
                'consommation: double',
                'enCours: boolean',
                'departLat: double',
                'departLng: double',
                'arriveeLat: double',
                'arriveeLng: double',
                'waypoints: String (JSON [[lat,lng]...])',
            ],
        },
    }

    boxes = {}  # name -> (cx, top_y, bottom_y)
    for name, info in entities.items():
        top_y, bottom_y = draw_class_box(
            ax, info['cx'], info['cy'], name, info['fields'],
            width=w, row_h=row_h, header_color=info['color']
        )
        boxes[name] = (info['cx'], top_y, bottom_y)

    def rel(name1, edge1, name2, edge2, label1='', label2='', rlabel='',
            color='#2c3e50', dashed=False):
        """Draw relationship line between two entity edges.
           edge: 'top','bottom','left','right'
        """
        cx1, top1, bot1 = boxes[name1]
        cx2, top2, bot2 = boxes[name2]

        def pt(name, edge):
            cx, top, bot = boxes[name]
            mid_y = (top + bot) / 2
            if edge == 'top':    return cx, top
            if edge == 'bottom': return cx, bot
            if edge == 'left':   return cx - w / 2, mid_y
            if edge == 'right':  return cx + w / 2, mid_y

        x1, y1 = pt(name1, edge1)
        x2, y2 = pt(name2, edge2)

        lstyle = 'dashed' if dashed else 'solid'
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='-', color=color, lw=1.8,
                                    linestyle=lstyle), zorder=3)
        # Multiplicity labels
        off = 0.18
        if label1:
            dx, dy = x2 - x1, y2 - y1
            norm = (dx**2 + dy**2) ** 0.5 or 1
            ax.text(x1 + dx * 0.10, y1 + dy * 0.10 + off,
                    label1, fontsize=8, color='#922b21', fontweight='bold',
                    ha='center', zorder=6)
        if label2:
            dx, dy = x2 - x1, y2 - y1
            ax.text(x1 + dx * 0.90, y1 + dy * 0.90 + off,
                    label2, fontsize=8, color='#922b21', fontweight='bold',
                    ha='center', zorder=6)
        if rlabel:
            mx, my = (x1 + x2) / 2, (y1 + y2) / 2
            ax.text(mx, my + 0.16, rlabel, fontsize=7.5, color='#555',
                    ha='center', style='italic', zorder=6,
                    bbox=dict(facecolor='white', edgecolor='none', alpha=0.7, pad=1))

    # Relationships
    rel('Chauffeur', 'right', 'Vehicule', 'left',
        label1='0..1', label2='0..1', rlabel='conducteurId', dashed=True)
    rel('Vehicule', 'bottom', 'Entretien', 'top',
        label1='1', label2='0..*', rlabel='vehiculeId')
    rel('Vehicule', 'bottom', 'Incident', 'top',
        label1='1', label2='0..*', rlabel='vehiculeId')
    rel('Vehicule', 'right', 'Trajet', 'left',
        label1='1', label2='0..*', rlabel='vehiculeId')
    rel('Chauffeur', 'bottom', 'Trajet', 'left',
        label1='1', label2='0..*', rlabel='chauffeurId')

    ax.set_title('Diagramme de classes — FleetTracking',
                 fontsize=16, fontweight='bold', pad=16, color='#1a252f')
    plt.tight_layout(pad=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  [2] Class diagram saved: {path}")


# ---------------------------------------------------------------------------
# DIAGRAM 3 — SEQUENCE: Marquer un entretien effectué
# ---------------------------------------------------------------------------

def draw_sequence_diagram_entretien(path):
    fig, ax = plt.subplots(figsize=(16, 11))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    participants = [
        ('Admin',               1.5,  '#1a5276'),
        ('Android App',         5.0,  '#117a65'),
        ('Backend\nSpring Boot',9.5,  '#6e2fa0'),
        ('PostgreSQL',         14.0,  '#922b21'),
    ]
    header_y = 10.4
    life_top  = 10.1
    life_bot  = 0.3

    # Participant boxes + lifelines
    for name, x, color in participants:
        box = FancyBboxPatch((x - 1.1, header_y - 0.32), 2.2, 0.65,
                             boxstyle="round,pad=0.08",
                             facecolor=color, edgecolor='#1a252f', lw=1.5, zorder=5)
        ax.add_patch(box)
        ax.text(x, header_y, name, ha='center', va='center', fontsize=9.5,
                fontweight='bold', color='white', zorder=6, multialignment='center')
        ax.plot([x, x], [life_top - 0.2, life_bot],
                color='#717d7e', lw=1.3, linestyle='--', dashes=(5, 4), zorder=2)

    pxs = [p[1] for p in participants]

    # Activation boxes
    act_specs = [
        (1, 9.85, 0.5),   # Android App
        (2, 7.8,  5.2),   # Backend
        (3, 7.5,  4.7),   # PostgreSQL
    ]
    for pidx, yt, yb in act_specs:
        px = pxs[pidx]
        ax.add_patch(FancyBboxPatch((px - 0.15, yb), 0.30, yt - yb,
                                    boxstyle="square,pad=0",
                                    facecolor='#d5d8dc', edgecolor='#717d7e',
                                    lw=1.2, zorder=3))

    # Arrow drawing helper
    def arrow(fi, ti, y, label, ret=False):
        x1, x2 = pxs[fi], pxs[ti]
        color = '#717d7e' if ret else '#1a252f'
        lstyle = 'dashed' if ret else 'solid'
        if fi == ti:
            # Self-arrow
            loop_x  = x1 + 0.15
            loop_dx = 0.75
            loop_dy = -0.30
            ax.annotate('', xy=(loop_x, y + loop_dy),
                        xytext=(loop_x, y),
                        arrowprops=dict(
                            arrowstyle='->',
                            color='#2c3e50', lw=1.5,
                            connectionstyle=f'arc3,rad=-0.7'),
                        zorder=4)
            ax.text(loop_x + 0.15, y + loop_dy / 2, label,
                    ha='left', va='center', fontsize=8.2, color='#1a252f', zorder=5,
                    bbox=dict(facecolor='white', edgecolor='none', alpha=0.75, pad=1))
        else:
            ax.annotate('', xy=(x2, y), xytext=(x1, y),
                        arrowprops=dict(arrowstyle='->', color=color, lw=1.5,
                                        linestyle=lstyle), zorder=4)
            mid_x = (x1 + x2) / 2
            ax.text(mid_x, y + 0.13, label, ha='center', va='bottom',
                    fontsize=8.2, color=color, zorder=5,
                    bbox=dict(facecolor='white', edgecolor='none', alpha=0.75, pad=1))

    # Step Y positions (top → bottom)
    ys = [9.55, 9.10, 8.65, 8.15, 7.65, 7.15, 6.65, 6.20, 5.65, 5.10, 4.55, 4.0, 3.45, 2.85]

    steps = [
        # (from, to, label, is_return)
        (0, 1, 'Presse "Marquer effectué" (bouton rouge)',          False),
        (1, 1, 'Affiche AlertDialog de confirmation',               False),
        (0, 1, 'Confirme',                                          False),
        (1, 2, 'PUT /api/entretiens/{id}/done (HTTP)',               False),
        (2, 3, "UPDATE entretien SET statut='effectue', date=today", False),
        (2, 2, 'Calcule retardJours (si date-based)',               False),
        (2, 2, 'Si récurrent: crée nouvel entretien',              False),
        (2, 3, 'INSERT INTO entretien (nouveau récurrent)',         False),
        (3, 2, '200 OK',                                            True),
        (2, 1, '200 OK + Entretien updated (JSON)',                  True),
        (1, 1, 'Invalide cache SQLite',                             False),
        (1, 2, 'GET /api/entretiens (rechargement)',                False),
        (2, 1, 'Liste entretiens mise à jour',                      True),
        (1, 0, 'Affiche liste rafraîchie (badge vert → disparu)',   True),
    ]

    for i, (fi, ti, label, ret) in enumerate(steps):
        y = ys[i]
        # Step number
        ax.text(0.22, y, str(i + 1), ha='center', va='center', fontsize=7.5,
                color='white', zorder=7,
                bbox=dict(boxstyle='circle,pad=0.18', facecolor='#2c3e50', edgecolor='none'))
        arrow(fi, ti, y, label, ret=ret)

    ax.set_title("Diagramme de séquence — Marquer un entretien effectué",
                 fontsize=13, fontweight='bold', pad=14, color='#1a252f')
    plt.tight_layout(pad=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  [3] Entretien sequence diagram saved: {path}")


# ---------------------------------------------------------------------------
# DIAGRAM 4 — SEQUENCE: Suivi de trajet GPS
# ---------------------------------------------------------------------------

def draw_sequence_diagram_gps(path):
    fig, ax = plt.subplots(figsize=(18, 13))
    ax.set_xlim(0, 18)
    ax.set_ylim(0, 13)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    participants = [
        ('Chauffeur',              1.2,  '#1a5276'),
        ('Android\n(TripManager)', 4.5,  '#117a65'),
        ('GPS\n(FusedLocation)',   8.0,  '#d35400'),
        ('Backend\nSpring Boot',  12.0,  '#6e2fa0'),
        ('Admin\n(CarteFragment)', 16.5, '#922b21'),
    ]
    header_y = 12.3
    life_top  = 12.0
    life_bot  = 0.35

    for name, x, color in participants:
        box = FancyBboxPatch((x - 1.05, header_y - 0.35), 2.1, 0.70,
                             boxstyle="round,pad=0.08",
                             facecolor=color, edgecolor='#1a252f', lw=1.5, zorder=5)
        ax.add_patch(box)
        ax.text(x, header_y, name, ha='center', va='center', fontsize=9,
                fontweight='bold', color='white', zorder=6, multialignment='center')
        ax.plot([x, x], [life_top - 0.1, life_bot],
                color='#717d7e', lw=1.3, linestyle='--', dashes=(5, 4), zorder=2)

    pxs = [p[1] for p in participants]

    # Activation boxes
    act_specs = [
        (1, 11.6, 0.5),   # Android active throughout
        (2, 8.9,  5.8),   # GPS during trip
        (3, 10.5, 5.0),   # Backend
    ]
    for pidx, yt, yb in act_specs:
        px = pxs[pidx]
        ax.add_patch(FancyBboxPatch((px - 0.15, yb), 0.30, yt - yb,
                                    boxstyle="square,pad=0",
                                    facecolor='#d5d8dc', edgecolor='#717d7e',
                                    lw=1.2, zorder=3))

    def arrow(fi, ti, y, label, ret=False):
        x1, x2 = pxs[fi], pxs[ti]
        color = '#717d7e' if ret else '#1a252f'
        lstyle = 'dashed' if ret else 'solid'
        if fi == ti:
            ax.annotate('', xy=(x1 + 0.15, y - 0.28),
                        xytext=(x1 + 0.15, y),
                        arrowprops=dict(
                            arrowstyle='->',
                            color='#2c3e50', lw=1.5,
                            connectionstyle='arc3,rad=-0.7'),
                        zorder=4)
            ax.text(x1 + 0.35, y - 0.14, label,
                    ha='left', va='center', fontsize=7.8, color='#1a252f', zorder=5,
                    bbox=dict(facecolor='white', edgecolor='none', alpha=0.75, pad=1))
        else:
            ax.annotate('', xy=(x2, y), xytext=(x1, y),
                        arrowprops=dict(arrowstyle='->', color=color, lw=1.5,
                                        linestyle=lstyle), zorder=4)
            mid_x = (x1 + x2) / 2
            ax.text(mid_x, y + 0.12, label, ha='center', va='bottom',
                    fontsize=7.8, color=color, zorder=5,
                    bbox=dict(facecolor='white', edgecolor='none', alpha=0.75, pad=1))

    # Y positions for steps
    # Pre-loop: steps 1-7 from top
    ys_pre  = [11.55, 11.05, 10.55, 10.05, 9.55, 9.05, 8.55]
    # Loop block: y range
    loop_top    = 8.25
    loop_bottom = 5.75
    ys_loop = [7.95, 7.40, 6.85, 6.30]
    # Post-loop: steps 12-20
    ys_post = [5.40, 4.90, 4.40, 3.90, 3.40, 2.90, 2.40, 1.90, 1.35]

    pre_steps = [
        (0, 1, 'Presse "Démarrer trajet"',                                    False),
        (1, 1, 'Vérifie permission ACCESS_FINE_LOCATION',                     False),
        (1, 1, 'TripManager.startTrip() — initialise waypoints=[]',           False),
        (1, 3, 'POST /api/trajets {vehiculeId, chauffeurId, heureDepart, enCours=true}', False),
        (3, 1, '201 Created {trajetId}',                                      True),
        (1, 2, 'requestLocationUpdates(interval=2s) — waypoints',             False),
        (1, 2, 'requestLocationUpdates(interval=30s) — position live',        False),
    ]
    loop_steps = [
        (2, 1, 'onLocationResult (2s) → waypoints.add([lat,lng])',           False),
        (2, 1, 'onLocationResult (30s) → pushPosition()',                    False),
        (1, 3, 'PUT /api/vehicules/{id} {lat, lng, vitesse}',                False),
        (3, 4, '(Admin poll) GET /api/vehicules → position mise à jour',     True),
    ]
    post_steps = [
        (0, 1, 'Presse "Arrêter trajet"',                                    False),
        (1, 1, 'TripManager.stopTrip()',                                      False),
        (1, 2, 'removeLocationUpdates() (both callbacks)',                    False),
        (1, 1, 'computeDistanceKm() via haversine sur waypoints',            False),
        (1, 1, 'serializeWaypoints() → JSON string',                         False),
        (1, 3, 'PUT /api/trajets/{id} {heureArrivee, distanceKm, duree, vitesseMoyenne, waypoints JSON, enCours=false}', False),
        (3, 1, '200 OK + Trajet complet',                                    True),
        (1, 0, '"Trajet arrêté" toast',                                      True),
        (4, 3, 'GET /api/trajets → voit trajet dans historique',             False),
    ]

    step_num = 1
    for i, (fi, ti, label, ret) in enumerate(pre_steps):
        y = ys_pre[i]
        ax.text(0.28, y, str(step_num), ha='center', va='center', fontsize=7,
                color='white', zorder=7,
                bbox=dict(boxstyle='circle,pad=0.15', facecolor='#2c3e50', edgecolor='none'))
        arrow(fi, ti, y, label, ret=ret)
        step_num += 1

    # Draw loop rectangle
    loop_lx = pxs[1] - 0.6
    loop_rx = pxs[4] + 0.7
    loop_rect = FancyBboxPatch((loop_lx, loop_bottom), loop_rx - loop_lx, loop_top - loop_bottom,
                               boxstyle="square,pad=0",
                               facecolor='#f9f9f9', edgecolor='#2471a3',
                               lw=2, linestyle='dashed', zorder=2, alpha=0.85)
    ax.add_patch(loop_rect)
    ax.text(loop_lx + 0.12, loop_top - 0.08, 'loop',
            ha='left', va='top', fontsize=9, fontweight='bold', color='#2471a3',
            bbox=dict(facecolor='#ddeeff', edgecolor='#2471a3', pad=3, lw=1))
    ax.text((loop_lx + loop_rx) / 2, loop_top - 0.12,
            'Boucle pendant le trajet',
            ha='center', va='top', fontsize=8.5, color='#2471a3', style='italic')

    for i, (fi, ti, label, ret) in enumerate(loop_steps):
        y = ys_loop[i]
        ax.text(0.28, y, str(step_num), ha='center', va='center', fontsize=7,
                color='white', zorder=7,
                bbox=dict(boxstyle='circle,pad=0.15', facecolor='#2471a3', edgecolor='none'))
        arrow(fi, ti, y, label, ret=ret)
        step_num += 1

    for i, (fi, ti, label, ret) in enumerate(post_steps):
        y = ys_post[i]
        ax.text(0.28, y, str(step_num), ha='center', va='center', fontsize=7,
                color='white', zorder=7,
                bbox=dict(boxstyle='circle,pad=0.15', facecolor='#2c3e50', edgecolor='none'))
        arrow(fi, ti, y, label, ret=ret)
        step_num += 1

    ax.set_title('Diagramme de séquence — Suivi de trajet GPS',
                 fontsize=13, fontweight='bold', pad=14, color='#1a252f')
    plt.tight_layout(pad=0.5)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    print(f"  [4] GPS sequence diagram saved: {path}")


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    tmp_dir = tempfile.mkdtemp()
    uc_path   = os.path.join(tmp_dir, 'use_case.png')
    cls_path  = os.path.join(tmp_dir, 'class_diag.png')
    seq1_path = os.path.join(tmp_dir, 'seq_entretien.png')
    seq2_path = os.path.join(tmp_dir, 'seq_gps.png')

    print("=" * 60)
    print("Generating 4 diagrams...")
    print("=" * 60)
    generate_use_case_diagram(uc_path)
    generate_class_diagram(cls_path)
    draw_sequence_diagram_entretien(seq1_path)
    draw_sequence_diagram_gps(seq2_path)

    print(f"\nOpening: {DOCX_PATH}")
    doc = Document(DOCX_PATH)
    print(f"  Total paragraphs before cleanup: {len(doc.paragraphs)}")

    print("\nRemoving old image paragraphs...")
    remove_image_paragraphs(doc)
    print(f"  Total paragraphs after cleanup: {len(doc.paragraphs)}")

    # ── Print all headings to help debug insertion ──
    print("\nDocument headings:")
    for i, p in enumerate(doc.paragraphs):
        if p.style.name.startswith('Heading'):
            try:
                print(f"  [{i:3d}] {p.style.name}: {p.text[:80]}")
            except Exception:
                print(f"  [{i:3d}] {p.style.name}: <encode error>")

    # ── Insert diagrams ──
    print("\nInserting diagrams...")

    # DIAGRAM 2 (Class) → after "3. Modèle de données" heading
    idx = find_paragraph_index(doc, '3.', style_contains='Heading')
    if idx is None:
        idx = find_paragraph_index(doc, 'mod')
    if idx is not None:
        print(f"  Inserting Class diagram after para {idx}: '{doc.paragraphs[idx].text[:60]}'")
        insert_image_after_paragraph(doc, idx, cls_path, width_inches=6.5,
                                     caption='Figure 2 — Diagramme de classes')
    else:
        print("  WARNING: Could not find section 3 heading for class diagram")

    # DIAGRAM 1 (Use Case) → after "4. Fonctionnalités principales" heading
    idx = find_paragraph_index(doc, '4.', style_contains='Heading')
    if idx is None:
        idx = find_paragraph_index(doc, 'fonctionnalit')
    if idx is not None:
        print(f"  Inserting Use Case diagram after para {idx}: '{doc.paragraphs[idx].text[:60]}'")
        insert_image_after_paragraph(doc, idx, uc_path, width_inches=6.5,
                                     caption="Figure 1 — Diagramme des cas d'utilisation")
    else:
        print("  WARNING: Could not find section 4 heading for use case diagram")

    # DIAGRAM 3 (Entretien sequence) → after "5. Système d'entretiens" heading
    idx = find_paragraph_index(doc, '5.', style_contains='Heading')
    if idx is None:
        idx = find_paragraph_index(doc, 'entretien', 'notif')
    if idx is not None:
        print(f"  Inserting Entretien seq diagram after para {idx}: '{doc.paragraphs[idx].text[:60]}'")
        insert_image_after_paragraph(doc, idx, seq1_path, width_inches=6.5,
                                     caption='Figure 3 — Séquence: Marquer un entretien effectué')
    else:
        print("  WARNING: Could not find section 5 heading for entretien sequence diagram")

    # DIAGRAM 4 (GPS sequence) → immediately after diagram 3 (same section)
    # Re-find the section 5 heading, then the next paragraph after the images we just inserted
    idx = find_paragraph_index(doc, '5.', style_contains='Heading')
    if idx is None:
        idx = find_paragraph_index(doc, 'entretien')
    if idx is not None:
        # Insert after the caption of diagram 3, which is now at idx+2
        insert_after_idx = min(idx + 2, len(doc.paragraphs) - 1)
        print(f"  Inserting GPS seq diagram after para {insert_after_idx}")
        insert_image_after_paragraph(doc, insert_after_idx, seq2_path, width_inches=6.5,
                                     caption='Figure 4 — Séquence: Suivi de trajet GPS')
    else:
        print("  WARNING: Could not find anchor paragraph for GPS sequence diagram")

    print(f"\nSaving document...")
    doc.save(DOCX_PATH)
    size_kb = os.path.getsize(DOCX_PATH) / 1024
    print(f"  Saved. File size: {size_kb:.0f} KB")

    # Count images
    img_count = sum(1 for p in Document(DOCX_PATH).paragraphs if paragraph_has_image(p))
    print(f"  Embedded image paragraphs in saved file: {img_count}")

    # Cleanup
    for f in [uc_path, cls_path, seq1_path, seq2_path]:
        try:
            os.remove(f)
        except Exception:
            pass
    try:
        os.rmdir(tmp_dir)
    except Exception:
        pass
    print("\nDone! Temp files cleaned up.")
    print("=" * 60)


if __name__ == '__main__':
    main()
